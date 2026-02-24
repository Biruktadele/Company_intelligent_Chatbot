import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import requests
from urllib.parse import urlparse
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
try:
    import textract
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False
from dataclasses import dataclass
from datetime import datetime
import hashlib

@dataclass
class DocumentMetadata:
    filename: str
    file_type: str
    size_bytes: int
    created_at: datetime
    modified_at: datetime
    source: str  # local, url, database
    checksum: Optional[str] = None
@dataclass
class AcquiredDocument:
    """Represents an acquired document with content and metadata"""
    id : str
    content: str
    metadata: DocumentMetadata
    raw_content: Optional[bytes] = None
class DocumentAcquisition:
    """Handles document acquisition from various sources"""
    
    SUPPORTED_FORMATS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
        '.txt': 'txt',
        '.html': 'html',
        '.htm': 'html',
        '.md': 'markdown',
        '.csv': 'csv'
    }
    def __init__(self, logger: Optional[logging.Logger] = None):
        self.logger = logger or logging.getLogger(__name__)
        
    def acquire_from_directory(self, directory_path: str, recursive: bool = True) -> List[AcquiredDocument]:
        """Acquire documents from a local directory"""
        documents = []
        directory = Path(directory_path)
        
        if not directory.exists():
            self.logger.error(f"Directory does not exist: {directory_path}")
            return documents
            
        pattern = "**/*" if recursive else "*"
        for file_path in directory.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_FORMATS:
                try:
                    doc = self._process_local_file(file_path)
                    if doc:
                        documents.append(doc)
                except Exception as e:
                    self.logger.error(f"Error processing file {file_path}: {str(e)}")
                    
        return documents
    def acquire_single_file(self, file_path: str) -> Optional[AcquiredDocument]:
        """Acquire a single local file"""
        path = Path(file_path)
        if not path.exists():
            self.logger.error(f"File does not exist: {file_path}")
            return None
            
        return self._process_local_file(path)
    def acquire_from_url(self, url: str) -> Optional[AcquiredDocument]:
        """Acquire document from a URL"""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            # Determine file type from URL or content-type
            parsed_url = urlparse(url)
            filename = os.path.basename(parsed_url.path) or "downloaded_document"
            
            content_type = response.headers.get('content-type', '').lower()
            if 'html' in content_type:
                file_type = 'html'
            elif 'pdf' in content_type:
                file_type = 'pdf'
            else:
                file_type = self._infer_file_type(filename)
                
            # Create metadata
            metadata = DocumentMetadata(
                filename=filename,
                file_type=file_type,
                size_bytes=len(response.content),
                created_at=datetime.now(),
                modified_at=datetime.now(),
                source='url'
            )
            
            # Extract content
            content = self._extract_content_from_bytes(response.content, file_type)
            
            return AcquiredDocument(
                id = hashlib.sha256(content.encode('utf-8')).hexdigest(),
                content=content,
                metadata=metadata,
                raw_content=response.content
            )
            
        except Exception as e:
            self.logger.error(f"Error acquiring document from URL {url}: {str(e)}")
            return None
    def _process_local_file(self, file_path: Path) -> Optional[AcquiredDocument]:
        """Process a single local file"""
        try:
            stat = file_path.stat()
            file_type = self._infer_file_type(file_path.name)
            
            metadata = DocumentMetadata(
                filename=file_path.name,
                file_type=file_type,
                size_bytes=stat.st_size,
                created_at=datetime.fromtimestamp(stat.st_ctime),
                modified_at=datetime.fromtimestamp(stat.st_mtime),
                source='local'
            )
            
            # Extract content based on file type
            if file_type == 'txt':
                content = self._extract_txt_content(file_path)
            elif file_type == 'pdf':
                content = self._extract_pdf_content(file_path)
            elif file_type == 'docx':
                content = self._extract_docx_content(file_path)
            elif file_type == 'html':
                content = self._extract_html_content(file_path)
            elif file_type == 'markdown':
                content = self._extract_markdown_content(file_path)
            else:
                # Use textract for other formats
                content = self._extract_with_textract(file_path)
                
            return AcquiredDocument(
                id = hashlib.sha256(content.encode('utf-8')).hexdigest(),
                content=content,
                metadata=metadata
                )
            
        except Exception as e:
            self.logger.error(f"Error processing file {file_path}: {str(e)}")
            return None
    def _infer_file_type(self, filename: str) -> str:
        """Infer file type from filename"""
        suffix = Path(filename).suffix.lower()
        return self.SUPPORTED_FORMATS.get(suffix, 'unknown')
    def _extract_txt_content(self, file_path: Path) -> str:
        """Extract content from a text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract text content from PDF file"""
        content = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                content.append(page.extract_text())
        return '\n'.join(content)

    def _extract_docx_content(self, file_path: Path) -> str:
        """Extract content from DOCX file"""
        doc = Document(file_path)
        content = []
        for paragraph in doc.paragraphs:
            content.append(paragraph.text)
        return '\n'.join(content)
    def _extract_html_content(self, file_path: Path) -> str:
        """Extract text content from HTML file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
        soup = BeautifulSoup(html_content, 'html.parser')
        return soup.get_text(separator='\n', strip=True)
    
    def _extract_markdown_content(self, file_path: Path) -> str:
        """Extract text content from Markdown file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    def _extract_content_from_bytes(self, content_bytes: bytes, file_type: str) -> str:
        """Extract content from bytes (for URL downloads)"""
        if file_type == 'html':
            soup = BeautifulSoup(content_bytes, 'html.parser')
            return soup.get_text(separator='\n', strip=True)
        else:
            # For other types, you might need to save to temp file first
            # This is a simplified implementation
            return content_bytes.decode('utf-8', errors='ignore')
    def _validate_content(self, content: str) -> str:
        """ Validate and clean extracted content"""
        if not content or len(content.strip()) < 10:
            return ""
        
        # Remove excessive whitespace
        content = ' '.join(content.split())
        
        # Basic content validation
        if len(content) < 50:
            self.logger.warning("Document content seems very short")
        
        return content
    
    def _extract_with_textract(self, file_path: Path) -> str:
        """Extract content using textract library"""
        if not TEXTRACT_AVAILABLE:
            self.logger.warning(f"Textract not available, cannot extract from {file_path}")
            return ""
        
        try:
            return textract.process(str(file_path)).decode('utf-8')
        except Exception as e:
            self.logger.error(f"Textract failed for {file_path}: {str(e)}")
            return ""
    
    
